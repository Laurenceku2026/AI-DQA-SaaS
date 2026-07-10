"""AI-DQA client DFMEA / report template fill helpers."""
from __future__ import annotations

MODULE_VERSION = "20260710"

import json
import os
import re
from datetime import datetime, timedelta
from io import BytesIO
from typing import Any, Callable, Dict, List, Optional, Tuple

try:
    from docx import Document
except Exception:  # pragma: no cover - docx/lxml may fail on some runtimes
    Document = None


def _load_workbook(source, **kwargs):
    """Lazy import so module load does not fail if openpyxl is temporarily unavailable."""
    from openpyxl import load_workbook

    return load_workbook(source, **kwargs)


def _extract_xls_outline(template_bytes: bytes, filename: str) -> str:
    """Read legacy .xls headers via xlrd when available."""
    try:
        import xlrd

        book = xlrd.open_workbook(file_contents=template_bytes)
        sheet = book.sheet_by_index(0)
        lines = [f"Sheet: {sheet.name}"]
        for row_idx in range(min(12, sheet.nrows)):
            values = []
            for col_idx in range(min(29, sheet.ncols)):
                value = sheet.cell_value(row_idx, col_idx)
                if value is not None and str(value).strip():
                    values.append(f"C{col_idx + 1}:{str(value).strip()}")
            if values:
                lines.append(f"R{row_idx + 1} " + " | ".join(values))
        return "\n".join(lines)
    except Exception:
        return f"Template file: {filename}"

TEMPLATE_EXTENSIONS = (".xlsx", ".xls", ".docx")
DFMEA_SHEET_NAME = "DFMEA标准表格"
DFMEA_DATA_START_ROW = 12

# 旧版 DFMEA（模板1）列映射
LEGACY_DFMEA_SHEET = "DFMEA"
LEGACY_DFMEA_DATA_START_ROW = 11
LEGACY_DFMEA_ROW_COLUMNS = {
    "focus_element": 2,
    "failure_mode": 3,
    "failure_effect": 4,
    "severity": 5,
    "failure_cause": 7,
    "occurrence": 8,
    "prevention_control": 9,
    "detection_control": 10,
    "detection": 11,
    "rpn": 12,
    "prevention_action": 13,
    "responsible_target": 14,
}
LEGACY_DFMEA_HEADER_CELLS = {
    "part_name": (5, 3),
    "design_owner": (5, 8),
    "project_model": (6, 3),
    "key_date": (6, 8),
    "prepared_by": (6, 14),
    "team": (7, 2),
    "fmea_date": (7, 13),
}

from dqa_template_profiles import (  # noqa: E402
    TEMPLATE_PROFILES,
    get_template_profile_label,
    profile_uses_deepseek_analysis,
    profile_uses_deepseek_fill,
    resolve_profile_template_filename,
)

DFMEA_FIELD_ALIASES: Dict[str, List[str]] = {
    "higher_level": ["higher_level", "上一级", "上一层级"],
    "focus_element": ["focus_element", "关注要素", "模块", "module"],
    "lower_level": ["lower_level", "下一层级", "下一低层级"],
    "higher_function": ["higher_function", "上一级功能"],
    "focus_function": ["focus_function", "关注要素功能"],
    "lower_function": ["lower_function", "下一级功能"],
    "failure_effect": ["failure_effect", "失效影响", "fe", "FE"],
    "severity": ["severity", "严重度", "s", "S"],
    "failure_mode": ["failure_mode", "失效模式", "fm", "FM", "关注要素的失效模式"],
    "failure_cause": ["failure_cause", "失效原因", "原因", "fc", "FC"],
    "prevention_control": ["prevention_control", "预防控制", "pc", "PC"],
    "occurrence": ["occurrence", "发生度", "o", "O"],
    "detection_control": ["detection_control", "探测控制", "dc", "DC"],
    "detection": ["detection", "探测度", "d", "D"],
    "action_priority": ["action_priority", "ap", "AP"],
    "prevention_action": ["prevention_action", "预防措施"],
    "detection_action": ["detection_action", "探测措施"],
    "responsible": ["responsible", "责任人", "负责人"],
    "target_date": ["target_date", "目标完成日期", "完成日期"],
}

DFMEA_HEADER_CELLS = {
    "project_name": (3, 17),
    "start_date": (4, 17),
    "revision_date": (5, 17),
    "customer_name": (5, 5),
    "model_year_project": (6, 5),
    "cross_function_team": (6, 9),
}

# DFMEA 步骤2-6 全字段列映射（1-based column index）
DFMEA_ROW_COLUMNS = {
    "higher_level": 4,
    "focus_element": 5,
    "lower_level": 6,
    "higher_function": 7,
    "focus_function": 8,
    "lower_function": 9,
    "failure_effect": 10,
    "severity": 11,
    "failure_mode": 12,
    "failure_cause": 13,
    "prevention_control": 14,
    "occurrence": 15,
    "detection_control": 16,
    "detection": 17,
    "action_priority": 18,
    "prevention_action": 19,
    "detection_action": 20,
    "responsible": 21,
    "target_date": 22,
}


def list_report_templates(app_key: str = "AI-DQA") -> List[str]:
    preferred = [
        "模板1-DFMEA旧版.xlsx",
        "模板2-DFMEA新版.xlsx",
    ]
    found: List[str] = []
    for name in preferred:
        try:
            resolve_template_path(name, app_key)
            if name not in found:
                found.append(name)
        except FileNotFoundError:
            continue
    here = os.path.dirname(os.path.abspath(__file__))
    folders = [
        os.path.join(here, "templates"),
        os.path.join(os.environ.get("DFSS_TEMPLATE_DIR", ""), app_key),
        os.path.join(r"C:\Users\Laurence\Technical\Project\SaaS\DFSS Report Template", app_key),
    ]
    for folder in folders:
        if not folder or not os.path.isdir(folder):
            continue
        for name in os.listdir(folder):
            if name.startswith("~$"):
                continue
            if os.path.splitext(name)[1].lower() in TEMPLATE_EXTENSIONS:
                if name not in found:
                    found.append(name)
    return sorted(found)


def template_mime_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".xls":
        return "application/vnd.ms-excel"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def resolve_template_path(filename: str, app_key: str = "AI-DQA") -> str:
    from dqa_template_profiles import normalize_template_filename

    filename = normalize_template_filename(filename)
    here = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(here, "templates", filename),
        os.path.join(os.environ.get("DFSS_TEMPLATE_DIR", ""), app_key, filename),
        os.path.join(r"C:\Users\Laurence\Technical\Project\SaaS\DFSS Report Template", app_key, filename),
    ]
    for path in candidates:
        if path and os.path.isfile(path):
            return path

    # Fallback: fuzzy match in templates/ (handles encoding / legacy names on Cloud)
    templates_dir = os.path.join(here, "templates")
    if os.path.isdir(templates_dir):
        markers = []
        if "旧版" in filename or "模板1" in filename:
            markers = ["旧版", "模板1"]
        elif "新版" in filename or "模板2" in filename:
            markers = ["新版", "模板2"]
        for name in os.listdir(templates_dir):
            if not name.endswith((".xlsx", ".xls")):
                continue
            if markers and any(m in name for m in markers):
                return os.path.join(templates_dir, name)

    raise FileNotFoundError(f"Template not found: {filename}")


def _clean_cell_text(text: str) -> str:
    return re.sub(r"\*\*", "", text or "").strip()


def _pick_field(row: Dict[str, Any], canonical: str) -> str:
    if not isinstance(row, dict):
        return ""
    aliases = DFMEA_FIELD_ALIASES.get(canonical, [canonical])
    for key in aliases:
        if key in row and row[key] is not None and str(row[key]).strip():
            return str(row[key]).strip()
    lowered = {str(k).lower(): v for k, v in row.items()}
    for key in aliases:
        val = lowered.get(str(key).lower())
        if val is not None and str(val).strip():
            return str(val).strip()
    return ""


def _normalize_dfmea_row(row: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(row, dict):
        return {}
    normalized = {field: _pick_field(row, field) for field in DFMEA_ROW_COLUMNS}
    if not normalized.get("failure_mode"):
        normalized["failure_mode"] = (
            _pick_field(row, "failure_mode")
            or normalized.get("focus_function", "")
            or normalized.get("focus_element", "")
        )
    if not normalized.get("failure_effect"):
        normalized["failure_effect"] = normalized.get("failure_mode", "")
    if not normalized.get("failure_cause") and normalized.get("lower_function"):
        pass
    return normalized


def _extract_risk_table_section(report_content: str) -> str:
    """Keep only the risk table section to reduce DeepSeek prompt size."""
    lines = report_content.splitlines()
    captured: List[str] = []
    started = False
    for line in lines:
        stripped = line.strip()
        if "|" in stripped and any(
            token in stripped.lower()
            for token in ["模块", "失效", "module", "failure", "严重", "severity", "rpn"]
        ):
            started = True
        if started:
            if stripped.startswith("|") or not stripped:
                captured.append(line)
            elif captured:
                break
    if captured:
        return "\n".join(captured)
    return report_content[:3500]


def _target_completion_date(fill_date: Optional[datetime] = None, days: int = 90) -> datetime:
    base = fill_date or datetime.now()
    return base + timedelta(days=days)


def _derive_row_extras(
    row: Dict[str, Any],
    lang: str,
    analyst_name: str = "",
    fill_date: Optional[datetime] = None,
) -> Dict[str, Any]:
    module = row.get("focus_element") or row.get("module", "")
    fm = row.get("failure_mode", "")
    cause = row.get("failure_cause") or row.get("cause", "")
    responsible = analyst_name or ("设计工程师" if lang == "zh" else "Design Engineer")
    target_dt = _target_completion_date(fill_date, 90)
    target_iso = target_dt.strftime("%Y-%m-%d")
    target_slash = target_dt.strftime("%Y/%m/%d")

    if not row.get("prevention_control"):
        row["prevention_control"] = (
            f"设计评审与{module}相关方案及规范符合性检查"
            if lang == "zh" and module
            else (f"Design review and compliance check for {module}" if module else (
                "设计评审与规范符合性检查" if lang == "zh" else "Design review and compliance check"
            ))
        )
    if not row.get("detection_control"):
        row["detection_control"] = (
            f"型式试验/可靠性测试验证{fm}" if lang == "zh" and fm
            else (f"Verification test for {fm}" if fm else (
                "型式试验与出厂检验" if lang == "zh" else "Type test and outgoing inspection"
            ))
        )
    if not row.get("prevention_action"):
        if lang == "zh":
            if cause and module:
                row["prevention_action"] = f"针对{cause}优化{module}设计"
            elif fm:
                row["prevention_action"] = f"优化设计以降低{fm}风险"
            else:
                row["prevention_action"] = "优化设计方案并更新设计规范"
        else:
            row["prevention_action"] = (
                f"Improve {module} design for {cause}" if cause and module
                else (f"Mitigate {fm} risk through design optimization" if fm else "Update design specification")
            )
    if not row.get("detection_action"):
        row["detection_action"] = row.get("detection_control", "")

    row["responsible"] = responsible
    row["target_date"] = target_iso
    row["responsible_target"] = f"  {responsible}{target_slash}"

    s, o, d = _to_int(row.get("severity")), _to_int(row.get("occurrence")), _to_int(row.get("detection"))
    if s and o and d:
        row["rpn"] = s * o * d
    return row


def _merge_rows_with_report(
    rows: List[Dict[str, Any]],
    product_name: str,
    product_desc: str,
    report_content: str,
    lang: str,
    analyst_name: str = "",
    fill_date: Optional[datetime] = None,
) -> List[Dict[str, Any]]:
    risks = parse_risks_from_markdown(report_content)
    fallback_rows = _simple_rows_from_report(
        product_name, product_desc, report_content, lang, analyst_name, fill_date
    )
    merged: List[Dict[str, Any]] = []
    source_rows = rows or fallback_rows
    for idx, row in enumerate(source_rows[:5]):
        normalized = _normalize_dfmea_row(row if isinstance(row, dict) else {})
        if risks and idx < len(risks):
            risk = risks[idx]
            if not normalized.get("failure_mode"):
                normalized["failure_mode"] = risk.get("failure_mode", "")
            if not normalized.get("failure_cause"):
                normalized["failure_cause"] = risk.get("cause", "")
            if not normalized.get("severity"):
                normalized["severity"] = risk.get("severity", "")
            if not normalized.get("occurrence"):
                normalized["occurrence"] = risk.get("occurrence", "")
            if not normalized.get("detection"):
                normalized["detection"] = risk.get("detection", "")
            if not normalized.get("focus_element"):
                normalized["focus_element"] = risk.get("module", "")
            for extra in ("prevention_control", "detection_control", "prevention_action"):
                if not normalized.get(extra) and risk.get(extra):
                    normalized[extra] = risk.get(extra)
        if idx < len(fallback_rows):
            fb = fallback_rows[idx]
            for field in (
                "failure_mode", "failure_effect", "failure_cause", "focus_element",
                "prevention_control", "detection_control", "prevention_action", "detection_action",
            ):
                if not normalized.get(field) and fb.get(field):
                    normalized[field] = fb[field]
        merged.append(_derive_row_extras(normalized, lang, analyst_name, fill_date))
    return merged or fallback_rows


def _parse_json_from_llm(text: Any) -> Any:
    if text is None:
        return None
    if not isinstance(text, str):
        text = str(text)
    if not text.strip():
        return None
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    for opener, closer in [("[", "]"), ("{", "}")]:
        start = cleaned.find(opener)
        end = cleaned.rfind(closer)
        if start != -1 and end > start:
            try:
                return json.loads(cleaned[start : end + 1])
            except json.JSONDecodeError:
                pass
    return None


def extract_template_outline(template_bytes: bytes, filename: str) -> str:
    """Backward-compatible wrapper."""
    from template_outline_utils import extract_template_outline as _extract

    return _extract(template_bytes, filename)


def build_template_guided_analysis_addon(template_outline: str, lang: str) -> str:
    if lang == "en":
        return f"""
=== Client DFMEA template requirements (Steps 2-6) ===
The analysis must be detailed enough to fill a DFMEA template with these fields per risk row:
- Step 2 structure: higher level, focus element, next lower level
- Step 3 function: function/requirements for each level
- Step 4 failure: failure effect (FE), severity (S), failure mode (FM), failure cause (FC)
- Step 5 risk: prevention control (PC), occurrence (O), detection control (DC), detection (D)
- Step 6 optimization: prevention actions, detection actions, owner, target date

Template outline:
{template_outline}
"""
    return f"""
=== 客户 DFMEA 模板要求（步骤2-6） ===
分析结果必须足够详细，可填入 DFMEA 模板每一行风险数据的以下字段：
- 步骤2 结构分析：上一层级、关注要素、下一低层级
- 步骤3 功能分析：各层级功能及要求
- 步骤4 失效分析：失效影响(FE)、严重度(S)、失效模式(FM)、失效原因(FC)
- 步骤5 风险分析：预防控制(PC)、发生度(O)、探测控制(DC)、探测度(D)
- 步骤6 优化：预防措施、探测措施、责任人、目标完成日期

模板结构摘要：
{template_outline}
"""


def generate_dfmea_rows_with_deepseek(
    template_outline: str,
    product_name: str,
    product_desc: str,
    report_content: str,
    lang: str,
    call_deepseek: Callable[[str, int], str],
    analyst_name: str = "",
) -> List[Dict[str, Any]]:
    """Use DeepSeek to map web report content into complete DFMEA rows."""
    risk_section = _extract_risk_table_section(report_content)
    prompt = f"""
你是资深可靠性工程师。根据以下风险分析表，生成 DFMEA 最多 5 行 JSON 数据。

产品：{product_name}
设计：{product_desc[:200]}

风险表：
{risk_section}

每个 JSON 对象必须包含且不可留空：
failure_mode（关注要素失效模式FM）, failure_effect（失效影响FE）, failure_cause（失效原因FC）,
focus_element, higher_level, lower_level, higher_function, focus_function, lower_function,
severity, occurrence, detection, prevention_control, detection_control,
prevention_action, detection_action, responsible, target_date

注意：failure_mode 与 failure_effect 必须分别填写，failure_mode 写失效模式，failure_effect 写后果。
severity/occurrence/detection 用 1-10 整数。target_date 格式 YYYY-MM-DD，应为填表日起约 90 天内完成。
语言：{'中文' if lang == 'zh' else 'English'}。只输出 JSON 数组。
"""
    raw = call_deepseek(prompt, 2800)
    parsed = _parse_json_from_llm(raw)
    if isinstance(parsed, list):
        return _merge_rows_with_report(
            parsed, product_name, product_desc, report_content, lang, analyst_name, datetime.now()
        )
    return []


def _normalize_header(header: str) -> str:
    h = _clean_cell_text(header).lower()
    mapping = {
        "模块": "module",
        "module": "module",
        "失效模式": "failure_mode",
        "failure mode": "failure_mode",
        "原因": "cause",
        "cause": "cause",
        "严重度": "severity",
        "severity": "severity",
        "发生度": "occurrence",
        "occurrence": "occurrence",
        "探测度": "detection",
        "detection": "detection",
        "预防控制": "prevention_control",
        "探测控制": "detection_control",
        "预防措施": "prevention_action",
        "建议措施": "prevention_action",
        "rpn": "rpn",
    }
    for key, value in mapping.items():
        if key in h:
            return value
    return h


def parse_risks_from_markdown(report_content: str) -> List[Dict[str, str]]:
    risks: List[Dict[str, str]] = []
    header_map: Optional[Dict[str, int]] = None

    for line in report_content.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            if header_map and risks:
                break
            continue

        cells = [_clean_cell_text(c) for c in line.split("|")[1:-1]]
        if not cells:
            continue
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue

        normalized = [_normalize_header(c) for c in cells]
        if "module" in normalized and "failure_mode" in normalized:
            header_map = {name: idx for idx, name in enumerate(normalized)}
            continue

        if not header_map:
            continue

        def pick(field: str) -> str:
            idx = header_map.get(field)
            if idx is None or idx >= len(cells):
                return ""
            return cells[idx]

        module = pick("module")
        failure_mode = pick("failure_mode")
        if not module and not failure_mode:
            continue

        risks.append(
            {
                "module": module,
                "failure_mode": failure_mode,
                "cause": pick("cause"),
                "severity": pick("severity"),
                "occurrence": pick("occurrence"),
                "detection": pick("detection"),
                "rpn": pick("rpn"),
                "prevention_control": pick("prevention_control"),
                "detection_control": pick("detection_control"),
                "prevention_action": pick("prevention_action"),
            }
        )

    return risks[:10]


def _simple_rows_from_report(
    product_name: str,
    product_desc: str,
    report_content: str,
    lang: str,
    analyst_name: str = "",
    fill_date: Optional[datetime] = None,
) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for risk in parse_risks_from_markdown(report_content):
        rows.append(
            _derive_row_extras(
                {
                    "higher_level": product_name,
                    "focus_element": risk.get("module", ""),
                    "lower_level": product_desc[:60] if product_desc else "-",
                    "higher_function": product_desc[:80] if product_desc else product_name,
                    "focus_function": risk.get("module", ""),
                    "lower_function": product_desc[:80] if product_desc else "",
                    "failure_effect": risk.get("failure_mode", ""),
                    "severity": risk.get("severity", ""),
                    "failure_mode": risk.get("failure_mode", ""),
                    "failure_cause": risk.get("cause", ""),
                    "prevention_control": risk.get("prevention_control", ""),
                    "occurrence": risk.get("occurrence", ""),
                    "detection_control": risk.get("detection_control", ""),
                    "detection": risk.get("detection", ""),
                    "prevention_action": risk.get("prevention_action", ""),
                    "detection_action": "",
                    "responsible": "",
                    "target_date": "",
                },
                lang,
                analyst_name,
                fill_date,
            )
        )
    if not rows:
        rows.append(
            _derive_row_extras(
                {
                    "higher_level": product_name,
                    "focus_element": product_name,
                    "lower_level": product_desc[:60] if product_desc else "-",
                    "failure_mode": "待补充失效模式" if lang == "zh" else "Pending failure mode",
                    "failure_cause": product_desc[:120] if product_desc else "",
                },
                lang,
                analyst_name,
                fill_date,
            )
        )
    return rows[:5]


def _set_cell(ws, row: int, col: int, value) -> None:
    if value is None:
        return
    text = str(value).strip()
    if not text:
        return
    from openpyxl.cell.cell import MergedCell
    from openpyxl.styles import Font

    cell = ws.cell(row=row, column=col)
    if isinstance(cell, MergedCell):
        for merged in ws.merged_cells.ranges:
            if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
                row, col = merged.min_row, merged.min_col
                break
    cell = ws.cell(row=row, column=col, value=value if str(value).isdigit() else text)
    cell.font = Font(color="000000")


def _set_header_cell(ws, row: int, col: int, value, append: bool = False) -> None:
    if value is None or not str(value).strip():
        return
    from openpyxl.cell.cell import MergedCell
    from openpyxl.styles import Font

    cell = ws.cell(row=row, column=col)
    if isinstance(cell, MergedCell):
        for merged in ws.merged_cells.ranges:
            if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
                row, col = merged.min_row, merged.min_col
                break
    if append:
        existing = ws.cell(row=row, column=col).value
        text = str(value).strip()
        if existing and text in str(existing):
            return
        combined = f"{existing}   {text}" if existing else text
    else:
        combined = str(value).strip()
    cell = ws.cell(row=row, column=col, value=combined)
    cell.font = Font(color="000000")


def _to_int(value) -> Optional[int]:
    if value is None:
        return None
    m = re.search(r"\d+", str(value))
    return int(m.group()) if m else None


def _is_legacy_dfmea_workbook(wb, template_filename: str = "") -> bool:
    if template_filename and ("旧版" in template_filename or "模板1" in template_filename):
        return True
    return DFMEA_SHEET_NAME not in wb.sheetnames and LEGACY_DFMEA_SHEET in wb.sheetnames


def _fill_legacy_dfmea_rows(ws, rows: List[Dict[str, Any]]) -> None:
    int_fields = {"severity", "occurrence", "detection", "rpn"}
    for idx, row_data in enumerate(rows[:10]):
        row = LEGACY_DFMEA_DATA_START_ROW + idx
        for field, col in LEGACY_DFMEA_ROW_COLUMNS.items():
            value = row_data.get(field)
            if field == "failure_effect" and not value:
                value = row_data.get("failure_mode", "")
            if field in int_fields:
                _set_cell(ws, row, col, _to_int(value))
            else:
                _set_cell(ws, row, col, value)


def fill_legacy_dfmea_workbook(
    wb,
    product_name: str,
    product_desc: str,
    report_content: str,
    analyst_name: str,
    analyst_title: str,
    lang: str,
) -> BytesIO:
    ws = wb[LEGACY_DFMEA_SHEET] if LEGACY_DFMEA_SHEET in wb.sheetnames else wb[wb.sheetnames[0]]
    today_dt = datetime.now()
    today = today_dt.strftime("%Y-%m-%d")
    today_slash = today_dt.strftime("%Y/%m/%d")
    team = analyst_name or ("未填写" if lang == "zh" else "N/A")
    if analyst_title:
        team = f"{team} ({analyst_title})"
    responsible = analyst_name or ("未填写" if lang == "zh" else "N/A")

    _set_cell(ws, 5, 3, product_name)
    _set_cell(ws, 5, 8, responsible)
    _set_cell(ws, 6, 3, product_name)
    _set_cell(ws, 6, 8, today_slash)
    _set_cell(ws, 6, 14, responsible)
    _set_cell(ws, 7, 2, team)
    _set_header_cell(ws, 7, 13, today_slash, append=True)

    rows = _merge_rows_with_report(
        [], product_name, product_desc, report_content, lang, analyst_name, today_dt
    )
    _fill_legacy_dfmea_rows(ws, rows)
    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out


def _fill_dfmea_rows(ws, rows: List[Dict[str, Any]]) -> None:
    int_fields = {"severity", "occurrence", "detection", "action_priority"}
    for idx, row_data in enumerate(rows[:10]):
        row = DFMEA_DATA_START_ROW + idx
        for field, col in DFMEA_ROW_COLUMNS.items():
            value = row_data.get(field)
            if field in int_fields:
                _set_cell(ws, row, col, _to_int(value))
            else:
                _set_cell(ws, row, col, value)


def fill_dfmea_workbook(
    wb,
    product_name: str,
    product_desc: str,
    report_content: str,
    analyst_name: str,
    analyst_title: str,
    lang: str,
    template_outline: str = "",
    call_deepseek: Optional[Callable[[str, int], str]] = None,
    template_filename: str = "",
) -> BytesIO:
    if _is_legacy_dfmea_workbook(wb, template_filename):
        return fill_legacy_dfmea_workbook(
            wb, product_name, product_desc, report_content, analyst_name, analyst_title, lang
        )

    ws = wb[DFMEA_SHEET_NAME] if DFMEA_SHEET_NAME in wb.sheetnames else wb[wb.sheetnames[0]]
    today_dt = datetime.now()
    today = today_dt.strftime("%Y-%m-%d")
    team = analyst_name or ("未填写" if lang == "zh" else "N/A")
    if analyst_title:
        team = f"{team} ({analyst_title})"

    header_values = {
        "project_name": product_name,
        "start_date": today,
        "revision_date": today,
        "customer_name": product_name,
        "model_year_project": product_desc[:80] if product_desc else product_name,
        "cross_function_team": team,
    }
    for key, (row, col) in DFMEA_HEADER_CELLS.items():
        _set_cell(ws, row, col, header_values.get(key))

    if call_deepseek:
        rows = generate_dfmea_rows_with_deepseek(
            template_outline,
            product_name,
            product_desc,
            report_content,
            lang,
            call_deepseek,
            analyst_name,
        )
        if not rows:
            rows = _merge_rows_with_report(
                [], product_name, product_desc, report_content, lang, analyst_name, today_dt
            )
    else:
        rows = _merge_rows_with_report(
            [], product_name, product_desc, report_content, lang, analyst_name, today_dt
        )

    _fill_dfmea_rows(ws, rows)
    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out


def fill_dfmea_template(
    product_name: str,
    product_desc: str,
    report_content: str,
    analyst_name: str = "",
    analyst_title: str = "",
    lang: str = "zh",
    template_filename: str = "模板2-DFMEA新版.xlsx",
    template_bytes: Optional[bytes] = None,
    template_outline: str = "",
    call_deepseek: Optional[Callable[[str, int], str]] = None,
) -> BytesIO:
    if template_bytes:
        wb = _load_workbook(BytesIO(template_bytes))
    else:
        wb = _load_workbook(resolve_template_path(template_filename, "AI-DQA"))
    if not template_outline and template_bytes:
        template_outline = extract_template_outline(template_bytes, template_filename)
    return fill_dfmea_workbook(
        wb,
        product_name,
        product_desc,
        report_content,
        analyst_name,
        analyst_title,
        lang,
        template_outline=template_outline,
        call_deepseek=call_deepseek,
        template_filename=template_filename,
    )


def fill_word_template_with_deepseek(
    template_bytes: bytes,
    product_name: str,
    product_desc: str,
    report_content: str,
    template_outline: str,
    lang: str,
    call_deepseek: Callable[[str, int], str],
) -> BytesIO:
    if Document is None:
        raise ValueError("python-docx is required for Word template export.")

    prompt = f"""
你是文档自动填表助手。根据分析报告，为 Word 模板生成键值对填表数据。

产品名称：{product_name}
设计描述：{product_desc}
模板结构：
{template_outline}

网页分析报告：
{report_content}

请只输出 JSON 对象，键为模板中出现的字段名/标签，值为要填写的内容。使用{'中文' if lang == 'zh' else 'English'}。
"""
    raw = call_deepseek(prompt, 4000)
    mapping = _parse_json_from_llm(raw)
    if not isinstance(mapping, dict):
        mapping = {
            "产品名称": product_name,
            "项目名称": product_name,
            "设计描述": product_desc,
            "报告内容": report_content,
        }

    doc = Document(BytesIO(template_bytes))
    replacements = {str(k): str(v) for k, v in mapping.items() if v is not None}

    def replace_in_text(text: str) -> str:
        updated = text
        for key, value in replacements.items():
            updated = updated.replace(f"{{{{{key}}}}}", value)
            updated = updated.replace(f"【{key}】", value)
            if key in updated and len(key) <= 20:
                updated = updated.replace(f"{key}：", f"{key}：{value}")
                updated = updated.replace(f"{key}:", f"{key}:{value}")
        return updated

    for para in doc.paragraphs:
        if para.text.strip():
            para.text = replace_in_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = replace_in_text(cell.text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def export_report_template(
    product_name: str,
    product_desc: str,
    report_content: str,
    analyst_name: str = "",
    analyst_title: str = "",
    lang: str = "zh",
    template_filename: Optional[str] = None,
    template_bytes: Optional[bytes] = None,
    template_outline: str = "",
    call_deepseek: Optional[Callable[[str, int], str]] = None,
) -> Tuple[BytesIO, str]:
    filename = template_filename or "模板2-DFMEA新版.xlsx"
    ext = os.path.splitext(filename)[1].lower()

    if template_bytes is None and template_filename:
        try:
            with open(resolve_template_path(template_filename, "AI-DQA"), "rb") as f:
                template_bytes = f.read()
        except FileNotFoundError:
            template_bytes = None

    if not template_outline and template_bytes:
        template_outline = extract_template_outline(template_bytes, filename)

    if ext == ".docx":
        if not template_bytes or not call_deepseek:
            raise ValueError("Word 模板需要上传文件并启用 DeepSeek 自动填表。" if lang == "zh" else "Word template requires upload and DeepSeek fill.")
        data = fill_word_template_with_deepseek(
            template_bytes,
            product_name,
            product_desc,
            report_content,
            template_outline,
            lang,
            call_deepseek,
        )
        return data, template_mime_type(filename)

    if ext in (".xlsx", ".xls"):
        data = fill_dfmea_template(
            product_name=product_name,
            product_desc=product_desc,
            report_content=report_content,
            analyst_name=analyst_name,
            analyst_title=analyst_title,
            lang=lang,
            template_filename=filename,
            template_bytes=template_bytes,
            template_outline=template_outline,
            call_deepseek=call_deepseek,
        )
        return data, template_mime_type(filename)

    raise ValueError("仅支持 Excel (.xlsx/.xls) 或 Word (.docx) 模板。" if lang == "zh" else "Only Excel or Word templates are supported.")
