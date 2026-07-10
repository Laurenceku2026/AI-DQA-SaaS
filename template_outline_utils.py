"""Extract template structure summaries for AI-DQA analysis guidance."""
from __future__ import annotations

import os
from io import BytesIO
from typing import Any, Union

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

DFMEA_SHEET_NAME = "DFMEA标准表格"
BufferLike = Union[bytes, bytearray, memoryview]


def _as_bytes(data: Any) -> bytes:
    if data is None:
        return b""
    if isinstance(data, (bytes, bytearray)):
        return bytes(data)
    if isinstance(data, memoryview):
        return data.tobytes()
    return bytes(data)


def _safe_name(filename: Any) -> str:
    if filename is None:
        return "template.xlsx"
    return str(filename)


def _load_workbook(source, **kwargs):
    from openpyxl import load_workbook

    return load_workbook(source, **kwargs)


def _extract_xls_outline(template_bytes: bytes, filename: str) -> str:
    try:
        import xlrd

        book = xlrd.open_workbook(file_contents=template_bytes)
        sheet = book.sheet_by_index(0)
        lines = [f"Sheet: {sheet.name}"]
        for row_idx in range(min(12, sheet.nrows)):
            values = []
            for col_idx in range(min(29, sheet.ncols)):
                value = sheet.cell_value(row_idx, col_idx)
                if value is not None and str(value).strip():
                    values.append(f"C{col_idx + 1}:{str(value).strip()}")
            if values:
                lines.append(f"R{row_idx + 1} " + " | ".join(values))
        return "\n".join(lines)
    except Exception:
        return f"Template file: {filename}"


def _extract_xlsx_outline(template_bytes: bytes, filename: str) -> str:
    wb = _load_workbook(BytesIO(template_bytes), read_only=True, data_only=True)
    try:
        ws = wb[DFMEA_SHEET_NAME] if DFMEA_SHEET_NAME in wb.sheetnames else wb[wb.sheetnames[0]]
        lines = [f"Sheet: {ws.title}"]
        for row_idx in range(1, 13):
            values = []
            for col_idx in range(1, 30):
                value = ws.cell(row=row_idx, column=col_idx).value
                if value is not None and str(value).strip():
                    values.append(f"C{col_idx}:{str(value).strip()}")
            if values:
                lines.append(f"R{row_idx} " + " | ".join(values))
        return "\n".join(lines)
    finally:
        wb.close()


def _extract_docx_outline(template_bytes: bytes, filename: str) -> str:
    if Document is None:
        return f"Template file: {filename}"
    doc = Document(BytesIO(template_bytes))
    lines = ["Word template paragraphs:"]
    for idx, para in enumerate(doc.paragraphs[:80], start=1):
        text = (getattr(para, "text", None) or "").strip()
        if text:
            lines.append(f"P{idx}: {text}")
    for t_idx, table in enumerate(doc.tables[:5], start=1):
        lines.append(f"Table{t_idx}:")
        for r_idx, row in enumerate(table.rows[:8]):
            cells = [(getattr(cell, "text", None) or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(f"  R{r_idx + 1}: " + " | ".join(cells))
    return "\n".join(lines)


def extract_template_outline(template_bytes: BufferLike, filename: str) -> str:
    """Extract template headers / placeholders for DeepSeek planning."""
    safe_filename = _safe_name(filename)
    try:
        raw = _as_bytes(template_bytes)
        if not raw:
            return f"Template file: {safe_filename}"

        ext = os.path.splitext(safe_filename)[1].lower()
        if ext == ".xls":
            return _extract_xls_outline(raw, safe_filename)
        if ext == ".xlsx":
            return _extract_xlsx_outline(raw, safe_filename)
        if ext == ".docx":
            return _extract_docx_outline(raw, safe_filename)
        return f"Template file: {safe_filename}"
    except Exception:
        return f"Template file: {safe_filename}"
